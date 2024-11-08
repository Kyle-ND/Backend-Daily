from docx import Document
import os


names_file_path = "Names/invited_names.txt"
template_letter_path = "letter/starting_letter.docx"  # input letter template
output_directory = "output/ReadyToSend"  # output folder


with open(names_file_path, "r") as file:
    invited_names = file.readlines()
letter_template = Document(template_letter_path)


for invited_name in invited_names:
    name = invited_name.strip()  

    personalized_letter = Document()
    for paragraph in letter_template.paragraphs:
        new_paragraph = personalized_letter.add_paragraph(paragraph.text.replace("[name]", name))

    personalized_letter.save(f"{output_directory}/letter_for_{name}.docx")
